"""普通文档内容提取：PDF / Word(.docx) / Excel / Markdown / 纯文本。

每个解析器都是"尽力而为"：库缺失或解析失败时返回占位说明，绝不抛异常，
保证多模态对话在缺依赖时依然可用（附件元信息仍会保存与展示）。
"""
import io


def extract_file_text(name: str, data: bytes) -> str:
    """按扩展名提取文档文本。name: 文件名；data: 文件字节。"""
    ext = (name or "").rsplit(".", 1)[-1].lower() if "." in (name or "") else ""
    try:
        if ext in ("md", "markdown", "txt", "text", "csv", "log"):
            return _decode_text(data)
        if ext == "pdf":
            return _extract_pdf(data)
        if ext == "docx":
            return _extract_docx(data)
        if ext in ("xlsx", "xlsm"):
            return _extract_xlsx(data)
        return f"[暂不支持解析该文件类型：{ext or 'unknown'}]"
    except Exception as exc:  # 解析失败不阻断对话
        return f"[文件解析失败：{exc}]"


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "gbk", "utf-16"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[未安装 pypdf 解析库，无法提取 PDF 文本]"
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip()) or "[PDF 未提取到文本（可能为扫描件）]"


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        return "[未安装 python-docx 解析库，无法提取 Word 文本]"
    document = docx.Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts) or "[Word 文档未提取到文本]"


def _extract_xlsx(data: bytes) -> str:
    try:
        import openpyxl
    except ImportError:
        return "[未安装 openpyxl 解析库，无法提取 Excel 文本]"
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[工作表：{ws.title}]\n" + "\n".join(rows))
    return "\n\n".join(parts) or "[Excel 未提取到文本]"
